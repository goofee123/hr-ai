"""Text extraction service for PDF and DOCX files."""

import io
from typing import Tuple

from PyPDF2 import PdfReader
from docx import Document


def extract_text_from_pdf(file_content: bytes) -> Tuple[str, dict]:
    """Extract text from a PDF file.

    Returns:
        Tuple of (extracted_text, metadata)
    """
    try:
        pdf_reader = PdfReader(io.BytesIO(file_content))
        text_parts = []

        for page_num, page in enumerate(pdf_reader.pages, 1):
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_parts.append(f"--- Page {page_num} ---\n{page_text}")

        full_text = "\n\n".join(text_parts)

        metadata = {
            "page_count": len(pdf_reader.pages),
            "extraction_method": "PyPDF2",
        }

        # Try to extract PDF metadata
        if pdf_reader.metadata:
            if pdf_reader.metadata.title:
                metadata["title"] = pdf_reader.metadata.title
            if pdf_reader.metadata.author:
                metadata["author"] = pdf_reader.metadata.author

        return full_text, metadata

    except Exception as e:
        return "", {"error": str(e), "extraction_method": "PyPDF2"}


def extract_text_from_docx(file_content: bytes) -> Tuple[str, dict]:
    """Extract text from a DOCX file.

    Returns:
        Tuple of (extracted_text, metadata)
    """
    try:
        doc = Document(io.BytesIO(file_content))

        # Extract paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Extract tables
        table_texts = []
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_text.append(row_text)
            if table_text:
                table_texts.append("\n".join(table_text))

        # Combine all text
        all_text = "\n\n".join(paragraphs)
        if table_texts:
            all_text += "\n\n--- Tables ---\n" + "\n\n".join(table_texts)

        metadata = {
            "paragraph_count": len(paragraphs),
            "table_count": len(doc.tables),
            "extraction_method": "python-docx",
        }

        # Try to extract document properties
        core_props = doc.core_properties
        if core_props.title:
            metadata["title"] = core_props.title
        if core_props.author:
            metadata["author"] = core_props.author

        return all_text, metadata

    except Exception as e:
        return "", {"error": str(e), "extraction_method": "python-docx"}


def extract_text(file_content: bytes, mime_type: str) -> Tuple[str, dict]:
    """Extract text from a file based on its MIME type.

    Args:
        file_content: Raw file bytes
        mime_type: MIME type of the file

    Returns:
        Tuple of (extracted_text, metadata)
    """
    if mime_type == "application/pdf":
        return extract_text_from_pdf(file_content)
    elif mime_type in [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ]:
        return extract_text_from_docx(file_content)
    else:
        return "", {"error": f"Unsupported file type: {mime_type}"}
